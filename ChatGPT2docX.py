import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from io import BytesIO
import os
import zipfile
import tempfile
import gradio as gr

def load_json_with_fallback(file_path):
    encodings = ['utf-8', 'latin-1', 'iso-8859-1']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                return json.load(file)
        except (UnicodeDecodeError, json.JSONDecodeError):
            continue
    raise ValueError("Impossible de lire le fichier JSON avec les encodages disponibles.")

def map_images_to_identifiers(images_dir):
    image_paths = {}

    if os.path.exists(images_dir):
        for file in os.listdir(images_dir):
            full_path = os.path.join(images_dir, file)
            if os.path.isfile(full_path):
                image_paths[file] = full_path

    print(f"Images mappées : {list(image_paths.keys())}")
    return image_paths

def find_image_for_asset(asset_pointer, image_paths):
    for file_name, path in image_paths.items():
        if file_name.startswith(asset_pointer):
            return path
    return None

def get_image_paths_from_zip(zip_file, extract_path):
    with zipfile.ZipFile(zip_file, 'r') as zip_ref:
        zip_ref.extractall(extract_path)

    image_dir = os.path.join(extract_path, 'Dalle-generations')
    image_paths = map_images_to_identifiers(image_dir)
    return image_paths

def add_styled_paragraph(doc, text, style=None, bold=False, underline=False, color=None, align=None):
    paragraph = doc.add_paragraph(text)
    run = paragraph.runs[0]

    if style:
        paragraph.style = style
    if bold:
        run.bold = True
    if underline:
        run.underline = True
    if color:
        run.font.color.rgb = color
    if align:
        paragraph.alignment = align

    paragraph.space_before = Pt(6)
    paragraph.space_after = Pt(6)

    return paragraph

def create_conversation_doc(conversation, images_path, output_dir):
    doc = Document()

    add_styled_paragraph(
        doc,
        conversation['title'],
        style='Title',
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER
    )
    doc.add_paragraph()  # Ajout d'un espacement sous le titre

    for node_id, node_data in conversation.get('mapping', {}).items():
        message = node_data.get('message')
        if not message:
            continue

        author = message['author']['role']
        parts = message['content'].get('parts', [])

        for part in parts:
            if isinstance(part, str) and part.strip():
                if author == 'user':
                    add_styled_paragraph(
                        doc,
                        f"Utilisateur : {part.strip()}",
                        bold=True,
                        underline=True,
                        color=RGBColor(0, 102, 204)
                    )
                else:
                    add_styled_paragraph(
                        doc,
                        f"Assistant : {part.strip()}"
                    )
            elif isinstance(part, dict) and part.get('content_type') == 'image_asset_pointer':
                asset_pointer = part.get('asset_pointer').split('/')[-1]
                image_file_path = find_image_for_asset(asset_pointer, images_path)

                if image_file_path and os.path.exists(image_file_path):
                    try:
                        with Image.open(image_file_path) as img:
                            img_stream = BytesIO()
                            img.save(img_stream, format='PNG')
                            img_stream.seek(0)

                            doc.add_paragraph()
                            doc.add_picture(img_stream, width=Inches(4))
                            last_paragraph = doc.paragraphs[-1]
                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            add_styled_paragraph(
                                doc,
                                f"(Image : {os.path.basename(image_file_path)})",
                                style='Normal',
                                align=WD_ALIGN_PARAGRAPH.CENTER
                            )
                    except Exception as e:
                        add_styled_paragraph(doc, f"(Erreur lors de l'insertion de l'image : {e})")
                else:
                    print(f"Image manquante pour : {asset_pointer}")

    output_file_path = os.path.join(output_dir, f"{conversation['title'].replace(' ', '_')}.docx")
    doc.save(output_file_path)
    return output_file_path

def generate_conversations_zip(zip_file):
    # Extraction des fichiers depuis le ZIP
    extract_path = tempfile.mkdtemp()
    image_paths = get_image_paths_from_zip(zip_file.name, extract_path)  # Utilisation correcte de `image_paths`
    json_file_path = os.path.join(extract_path, 'conversations.json')

    # Charger les données JSON
    conversations_data = load_json_with_fallback(json_file_path)

    # Génération des fichiers DOCX
    docx_temp_dir = tempfile.mkdtemp()
    docx_files = []

    for conversation in conversations_data:
        output_file = create_conversation_doc(conversation, image_paths, docx_temp_dir)  # Transmettre `image_paths`
        docx_files.append(output_file)

    # Création du fichier ZIP final
    zip_output_path = os.path.join(docx_temp_dir, "conversations_output.zip")
    with zipfile.ZipFile(zip_output_path, 'w') as zipf:
        for docx_file in docx_files:
            zipf.write(docx_file, os.path.basename(docx_file))

    return zip_output_path

def main_interface(zip_file):
    zip_file_output = generate_conversations_zip(zip_file)
    return zip_file_output

interface = gr.Interface(
    fn=main_interface,
    inputs=gr.File(label="Téléchargez votre fichier ZIP contenant les conversations et images"),
    outputs=gr.File(label="Téléchargez le fichier ZIP avec les fichiers DOCX générés"),
    title="Extraction de conversations ChatGPT avec style amélioré",
    description="""
### Instructions d'utilisation :

1. Téléchargez le fichier ZIP de sauvegarde de ChatGPT, contenant :
   - Un fichier `conversations.json`.
   - Un dossier `Dalle-generations` avec les images associées (si présentes).

2. Importez ce fichier ZIP ici.

3. Vous recevrez un ZIP contenant des fichiers DOCX :
   - Un fichier par conversation.
   - Les images insérées directement dans les documents.

---

### How to use:

1. Upload the ChatGPT backup ZIP file, containing:
   - A `conversations.json` file.
   - A `Dalle-generations` folder with associated images (if any).

2. Upload this ZIP here.

3. You will receive a ZIP with DOCX files:
   - One file per conversation.
   - Images inserted directly into the documents.
""",
)

interface.launch()
